"""
Handles loading documents from various formats:
- PDF (via PyMuPDF)
- TXT
- Markdown
- DOCX

Returns a consistent list of Document objects.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from src.utils.logger import get_logger

log = get_logger(__name__)


@dataclass
class Document:
    """Represents a loaded document with content and metadata."""

    content: str
    source: str
    page: Optional[int] = None
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        # Normalize whitespace
        self.content = re.sub(r"\s+", " ", self.content).strip()


class DocumentLoader:
    """
    Loads documents from a directory or single file.
    Supports: .pdf, .txt, .md, .docx
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def load_directory(self, directory: str | Path) -> List[Document]:
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        documents = []
        files = [
            f for f in directory.rglob("*")
            if f.is_file() and f.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]

        if not files:
            log.warning(f"No supported documents found in {directory}")
            return []

        log.info(f"Found {len(files)} documents in {directory}")

        for file_path in files:
            try:
                docs = self.load_file(file_path)
                documents.extend(docs)
                log.info(f"  ✓ Loaded: {file_path.name} ({len(docs)} pages/sections)")
            except Exception as e:
                log.error(f"  ✗ Failed to load {file_path.name}: {e}")

        log.info(f"Total documents loaded: {len(documents)}")
        return documents

    def load_file(self, file_path: str | Path) -> List[Document]:
        file_path = Path(file_path)
        ext = file_path.suffix.lower()

        loaders = {
            ".pdf": self._load_pdf,
            ".txt": self._load_text,
            ".md": self._load_text,
            ".docx": self._load_docx,
        }

        loader = loaders.get(ext)
        if not loader:
            raise ValueError(f"Unsupported file type: {ext}")

        return loader(file_path)

    def _load_pdf(self, file_path: Path) -> List[Document]:
        """Load PDF using PyMuPDF for fast, accurate extraction."""
        try:
            import fitz  
        except ImportError:
            raise ImportError("Install PyMuPDF: pip install pymupdf")

        documents = []
        with fitz.open(str(file_path)) as pdf:
            for page_num, page in enumerate(pdf, start=1):
                text = page.get_text("text")
                if text.strip():
                    documents.append(
                        Document(
                            content=text,
                            source=file_path.name,
                            page=page_num,
                            metadata={
                                "file_path": str(file_path),
                                "total_pages": len(pdf),
                                "format": "pdf",
                            },
                        )
                    )
        return documents

    def _load_text(self, file_path: Path) -> List[Document]:
        """Load plain text or markdown files."""
        content = file_path.read_text(encoding="utf-8", errors="replace")
        if not content.strip():
            return []

        return [
            Document(
                content=content,
                source=file_path.name,
                page=None,
                metadata={
                    "file_path": str(file_path),
                    "format": file_path.suffix.lstrip("."),
                },
            )
        ]

    def _load_docx(self, file_path: Path) -> List[Document]:
        """Load DOCX files using python-docx."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)

        if not content.strip():
            return []

        return [
            Document(
                content=content,
                source=file_path.name,
                metadata={"file_path": str(file_path), "format": "docx"},
            )
        ]